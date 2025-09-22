#!/usr/bin/env python3
# Generates sample PDFs, XLSX, CSV, and DOCX for Apache Tika testing.

import os, csv, sys, subprocess, importlib, datetime
OUTDIR = "tika_samples"

def ensure(pkg, import_name=None):
    try:
        importlib.import_module(import_name or pkg)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])

# deps
ensure("openpyxl")
ensure("python-docx", "docx")
ensure("reportlab")
ensure("Pillow", "PIL")

# std imports after install
from openpyxl import Workbook
from openpyxl.styles import Font
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from PIL import Image, ImageDraw, ImageFont

os.makedirs(OUTDIR, exist_ok=True)

def create_png_logo(path):
    img = Image.new("RGB", (400, 200), (240, 240, 240))
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("DejaVuSans-Bold.ttf", 64)
    except Exception:
        font = ImageFont.load_default()
    d.rectangle([10, 10, 390, 190], outline=(0, 0, 0), width=3)
    d.text((40, 70), "TIKA", font=font, fill=(0, 0, 0))
    img.save(path)

def create_csv(path):
    rows = [
        ["id","name","notes","amount","date"],
        [1,"Alice",'He said "hello"',123.45,"2025-09-22"],
        [2,"Боб","multi-line\nnote",-50,"2024-01-15"],
        [3,"李四","emoji 😀",0,"2023-06-30"],
    ]
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f, quoting=csv.QUOTE_ALL)
        w.writerows(rows)

def create_xlsx(path):
    wb = Workbook()
    wb.properties.creator = "Tika Samples"
    wb.properties.title = "Tika XLSX Sample"
    wb.properties.description = "Workbook with data, formulas, dates, hyperlink, and merged cells."

    ws = wb.active
    ws.title = "Data"
    header = ["Item","Qty","Price","Total","Date"]
    ws.append(header)
    ws["A1"].font = Font(bold=True)
    ws["B1"].font = Font(bold=True)
    ws["C1"].font = Font(bold=True)
    ws["D1"].font = Font(bold=True)
    ws["E1"].font = Font(bold=True)

    ws.append(["Widget",3,9.99,"=B2*C2", datetime.date(2024,1,15)])
    ws.append(["Gadget",5,2.5,"=B3*C3", datetime.date(2023,6,30)])
    ws.append(["Thingamajig",1,199.95,"=B4*C4", datetime.date(2025,9,22)])
    ws["E2"].number_format = "yyyy-mm-dd"
    ws["E3"].number_format = "yyyy-mm-dd"
    ws["E4"].number_format = "yyyy-mm-dd"
    ws["D6"] = "=SUM(D2:D4)"
    ws["C6"] = "Grand Total"

    ws.merge_cells("A8:E8")
    ws["A8"] = "Merged cell example"

    ws["A10"] = "Apache Tika"
    ws["A10"].hyperlink = "https://tika.apache.org/"
    ws["A10"].style = "Hyperlink"

    uni = wb.create_sheet("Unicode")
    uni.append(["Language","Sample"])
    uni.append(["Chinese","你好，世界"])
    uni.append(["Russian","Привет, мир"])
    uni.append(["Emoji","😀🚀📄"])

    wb.save(path)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0000FF"); rPr.append(color)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_docx(path, img_path):
    doc = Document()
    doc.core_properties.title = "Tika DOCX Sample"
    doc.core_properties.author = "Tika Samples"
    doc.core_properties.subject = "Paragraphs, lists, table, image, hyperlink"
    doc.core_properties.keywords = "tika, docx, sample"

    doc.add_heading("Tika Test Document", level=1)
    p = doc.add_paragraph("This is a ")
    p.add_run("sample").bold = True
    p.add_run(" DOCX file with ")
    p.add_run("varied").italic = True
    p.add_run(" content.")
    p2 = doc.add_paragraph()
    add_hyperlink(p2, "Apache Tika website", "https://tika.apache.org/")

    for item in ["Text extraction", "Tables", "Images", "Unicode – 你好, Привет, 😀"]:
        doc.add_paragraph(item, style="List Bullet")

    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Shading"
    for r in range(3):
        for c in range(3):
            table.cell(r, c).text = f"R{r+1}C{c+1}"

    doc.add_picture(img_path, width=Inches(2))
    doc.save(path)

def create_pdf_text(path):
    c = canvas.Canvas(path, pagesize=LETTER)
    c.setTitle("Tika PDF: Text & Metadata")
    c.setAuthor("Tika Samples")
    c.setSubject("Simple text page with metadata")
    c.setKeywords("tika, pdf, sample")
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 720, "Tika PDF: Text & Metadata")
    c.setFont("Helvetica", 12)
    text = c.beginText(72, 690)
    text.textLines([
        "This PDF contains plain text and document metadata.",
        "Use it to test Tika's /tika or /meta endpoints.",
        "ASCII-only text to avoid font issues."
    ])
    c.drawText(text)
    c.showPage()
    c.save()

def create_pdf_image_table(path, img_path):
    c = canvas.Canvas(path, pagesize=LETTER)
    c.setTitle("Tika PDF: Image & Table")
    c.setAuthor("Tika Samples")
    c.setSubject("Bitmap image and drawn table")
    c.setKeywords("tika, pdf, image, table")

    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 720, "PDF with Image & Table")
    c.drawImage(img_path, 72, 540, width=2*inch, height=1*inch, preserveAspectRatio=True, mask='auto')

    data = [["Item","Qty","Price"],
            ["Widget",3,"9.99"],
            ["Gadget",5,"2.50"],
            ["Thing",1,"199.95"]]
    x0, y0 = 72, 500
    colw = [200, 100, 100]
    rowh = 22
    rows = len(data)
    cols = len(colw)

    # grid
    for i in range(rows+1):
        c.line(x0, y0 - i*rowh, x0 + sum(colw), y0 - i*rowh)
    for j in range(cols+1):
        c.line(x0 + sum(colw[:j]), y0, x0 + sum(colw[:j]), y0 - rows*rowh)

    # text
    c.setFont("Helvetica", 12)
    for r in range(rows):
        for j in range(cols):
            c.drawString(x0 + 6 + sum(colw[:j]), y0 - (r+1)*rowh + 6, str(data[r][j]))

    c.showPage()
    c.save()

def main():
    png = os.path.join(OUTDIR, "logo.png")
    create_png_logo(png)

    create_csv(os.path.join(OUTDIR, "sample.csv"))
    create_xlsx(os.path.join(OUTDIR, "sample.xlsx"))
    create_docx(os.path.join(OUTDIR, "sample.docx"), png)
    create_pdf_text(os.path.join(OUTDIR, "text_metadata.pdf"))
    create_pdf_image_table(os.path.join(OUTDIR, "image_table.pdf"), png)

    print(f"Done. Files in: {OUTDIR}")

if __name__ == "__main__":
    main()
